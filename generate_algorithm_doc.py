#!/usr/bin/env python3
"""Generate Word document explaining the campus route planner algorithms."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Campus Route Planner", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Algorithm Explanation in Terms of Data Structures and Algorithms"
    )
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 1. Introduction
    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "The Campus Route Planner is a navigation application that finds the shortest walking "
        "route between two locations on a campus map (e.g., from Dorm C to the Gym). The problem "
        "is modeled as a graph theory problem: campus locations and intersections are vertices, "
        "paths between them are edges, and the distance of each path is the edge weight. The "
        "shortest path is computed using Dijkstra's Algorithm, which relies on a priority queue "
        "implemented as a binary min-heap."
    )

    # 2. Problem Formulation
    add_heading(doc, "2. Problem Formulation", level=1)
    add_para(doc, "Formally, the routing problem can be stated as follows:")
    add_bullet(doc, "Given: A weighted, undirected graph G = (V, E) where V is the set of vertices (locations and intersections) and E is the set of edges (walkable paths).")
    add_bullet(doc, "Each edge (u, v) has a non-negative weight w(u, v) representing distance in meters.")
    add_bullet(doc, "Given: A source vertex s and a destination vertex t.")
    add_bullet(doc, "Find: A path from s to t that minimizes the total sum of edge weights.")

    add_para(
        doc,
        "This is the classic Single-Source Shortest Path (SSSP) problem. Because all edge "
        "weights are non-negative, Dijkstra's Algorithm is the appropriate choice. It is "
        "guaranteed to find the optimal (shortest) path."
    )

    # 3. Data Structures
    add_heading(doc, "3. Data Structures Used", level=1)

    add_heading(doc, "3.1 Graph (Adjacency List)", level=2)
    add_para(
        doc,
        "The campus map is stored as a graph using an adjacency list representation. "
        "This is the standard data structure for sparse graphs where the number of edges "
        "is much smaller than |V|²."
    )
    add_bullet(doc, "Vertices (nodes): Stored in a Map (hash table) keyed by node ID. Each node contains an id, label, coordinates (x, y), and type (location or intersection).")
    add_bullet(doc, "Edges: Stored in an adjacency list — a Map where each key is a node ID and each value is an array of neighbor objects { to, weight, id }.")
    add_bullet(doc, "Undirected edges: Each road is added in both directions with the same weight.")

    add_para(doc, "Why adjacency list?", bold=True)
    add_bullet(doc, "Space complexity: O(V + E) — efficient for sparse campus maps.")
    add_bullet(doc, "Neighbor lookup: O(degree(v)) — iterate only over adjacent nodes during Dijkstra's relaxation step.")
    add_bullet(doc, "Easy to modify: Road closures can be handled by skipping edges during traversal.")

    add_heading(doc, "3.2 Priority Queue (Binary Min-Heap)", level=2)
    add_para(
        doc,
        "Dijkstra's Algorithm requires repeatedly selecting the unvisited vertex with the "
        "smallest tentative distance. A priority queue (min-heap) supports this efficiently."
    )
    add_bullet(doc, "Structure: A complete binary tree stored in an array, where each parent has a priority less than or equal to its children.")
    add_bullet(doc, "enqueue(value, priority): Insert an element and bubble up — O(log n).")
    add_bullet(doc, "dequeue(): Remove and return the minimum-priority element and bubble down — O(log n).")
    add_bullet(doc, "The priority of each vertex in the queue equals its current shortest known distance from the source.")

    add_para(doc, "Why a min-heap?", bold=True)
    add_bullet(doc, "Without a priority queue, finding the minimum-distance vertex requires scanning all vertices — O(V) per iteration, leading to O(V²) total time.")
    add_bullet(doc, "With a min-heap, each extract-min is O(log V), improving total time to O((V + E) log V).")

    add_heading(doc, "3.3 Hash Maps (Maps) and Sets", level=2)
    add_bullet(doc, "dist Map: Stores the shortest known distance from the source to each vertex. Initialized to Infinity, except the source (0).")
    add_bullet(doc, "prev Map: Stores the predecessor of each vertex on the shortest path — used for path reconstruction.")
    add_bullet(doc, "prevEdge Map: Stores which edge was used to reach each vertex — used to highlight the route on the map.")
    add_bullet(doc, "closedEdges Set: Stores IDs of blocked edges (road closures). Lookup is O(1) during neighbor relaxation.")

    add_para(
        doc,
        "All Map and Set operations (get, set, has) run in average O(1) time, making them "
        "ideal for tracking algorithm state across vertices."
    )

    # 4. Dijkstra's Algorithm
    add_heading(doc, "4. Dijkstra's Algorithm", level=1)

    add_heading(doc, "4.1 Overview", level=2)
    add_para(
        doc,
        "Dijkstra's Algorithm, published by Edsger W. Dijkstra in 1959, solves the single-source "
        "shortest path problem for graphs with non-negative edge weights. It uses a greedy strategy: "
        "at each step, it permanently settles the closest unvisited vertex and relaxes its neighbors."
    )

    add_heading(doc, "4.2 Algorithm Steps", level=2)
    steps = [
        "Initialize dist[s] = 0 for source s; dist[v] = ∞ for all other vertices.",
        "Insert source s into the priority queue with priority 0.",
        "While the priority queue is not empty:",
        "    a. Dequeue the vertex u with the smallest priority (distance).",
        "    b. If u is the destination, stop early (optional optimization).",
        "    c. For each neighbor v of u:",
        "       i.   Skip v if the edge (u, v) is closed/blocked.",
        "       ii.  Compute alt = dist[u] + weight(u, v).",
        "       iii. If alt < dist[v], update dist[v] = alt, set prev[v] = u, and enqueue v with priority alt.",
        "If dist[destination] = ∞, no path exists. Otherwise, reconstruct the path by following prev[] from destination back to source.",
    ]
    for step in steps:
        add_para(doc, step)

    add_heading(doc, "4.3 Key Concept: Relaxation", level=2)
    add_para(
        doc,
        "Relaxation is the core operation of Dijkstra's Algorithm. When considering an edge (u, v) "
        "with weight w, we ask: 'Is the path to v through u shorter than the current best path to v?' "
        "If dist[u] + w < dist[v], we update dist[v] and record u as v's predecessor. This is called "
        "relaxing the edge (u, v)."
    )

    add_heading(doc, "4.4 Path Reconstruction", level=2)
    add_para(
        doc,
        "After the algorithm completes, the shortest path is reconstructed by starting at the "
        "destination and following the prev map backward to the source. The path is then reversed "
        "to obtain the correct start-to-end order. The prevEdge map records which edges were used, "
        "allowing the UI to highlight the route on the map."
    )

    # 5. Complexity Analysis
    add_heading(doc, "5. Time and Space Complexity", level=1)

    add_heading(doc, "5.1 Time Complexity", level=2)
    add_bullet(doc, "Priority queue operations: Each vertex may be enqueued multiple times, but each edge is relaxed at most once. With a binary min-heap: O((V + E) log V).")
    add_bullet(doc, "For our campus graph: V = 7 vertices, E = 10 edges — the algorithm runs in negligible time.")
    add_bullet(doc, "Comparison: Using an unsorted array to find the minimum would be O(V²), which is acceptable for small graphs but does not scale.")

    add_heading(doc, "5.2 Space Complexity", level=2)
    add_bullet(doc, "Graph (adjacency list): O(V + E)")
    add_bullet(doc, "dist, prev, prevEdge maps: O(V) each")
    add_bullet(doc, "Priority queue: O(V) in the worst case")
    add_bullet(doc, "Total: O(V + E)")

    # 6. Worked Example
    add_heading(doc, "6. Worked Example: Dorm C → Gym", level=1)
    add_para(
        doc,
        "Using the campus graph from the project, finding the shortest path from Dorm C to the Gym:"
    )

    add_para(doc, "Graph vertices:", bold=True)
    add_bullet(doc, "Locations: Library, Engineering Hall, Dorm C, Gym, Cafeteria")
    add_bullet(doc, "Intersections: Upper intersection, Lower intersection")

    add_para(doc, "Relevant edges and weights (meters):", bold=True)
    add_bullet(doc, "Dorm C ↔ Lower intersection: 180 m")
    add_bullet(doc, "Lower intersection ↔ Cafeteria: 200 m")
    add_bullet(doc, "Cafeteria ↔ Gym: 230 m")
    add_bullet(doc, "Lower intersection ↔ Gym: 500 m (alternative, longer route)")

    add_para(doc, "Algorithm trace (simplified):", bold=True)
    add_bullet(doc, "Start: dist[Dorm C] = 0; all others = ∞.")
    add_bullet(doc, "Dequeue Dorm C (dist = 0). Relax neighbors: Lower intersection gets dist = 180.")
    add_bullet(doc, "Dequeue Lower intersection (dist = 180). Relax: Cafeteria → 380, Gym → 680.")
    add_bullet(doc, "Dequeue Cafeteria (dist = 380). Relax: Gym → 380 + 230 = 610.")
    add_bullet(doc, "Dequeue Gym (dist = 610). Destination reached. Stop.")

    add_para(doc, "Result:", bold=True)
    add_bullet(doc, "Shortest path: Dorm C → Lower intersection → Cafeteria → Gym")
    add_bullet(doc, "Total distance: 180 + 200 + 230 = 610 meters")
    add_bullet(doc, "Estimated walk time: 610 ÷ 78 ≈ 8 minutes (at ~1.3 m/s walking speed)")

    add_para(
        doc,
        "Note: The direct path from Lower intersection to Gym (500 m) would give a total of "
        "180 + 500 = 680 m, which is longer than the route through the Cafeteria. Dijkstra's "
        "Algorithm correctly chooses the shorter path."
    )

    # 7. Edge Cases
    add_heading(doc, "7. Edge Cases and Extensions", level=1)

    add_heading(doc, "7.1 Road Closures", level=2)
    add_para(
        doc,
        "The application supports simulating blocked paths (e.g., construction, events). "
        "Closed edges are stored in a Set and skipped during the relaxation step. If the "
        "shortest path uses a closed edge, the algorithm automatically finds an alternate route."
    )
    add_para(doc, "Example: If Cafeteria ↔ Gym is closed, the route becomes Dorm C → Lower intersection → Gym (680 m).")

    add_heading(doc, "7.2 No Path Exists", level=2)
    add_para(
        doc,
        "If road closures disconnect the source from the destination, dist[destination] remains "
        "Infinity after the algorithm completes. The application returns null and displays an "
        "appropriate message to the user."
    )

    add_heading(doc, "7.3 Same Source and Destination", level=2)
    add_para(
        doc,
        "When start equals end, the algorithm returns immediately with distance 0 and a path "
        "containing only that vertex."
    )

    # 8. Why Not Other Algorithms?
    add_heading(doc, "8. Why Dijkstra's Algorithm?", level=1)
    add_bullet(doc, "BFS (Breadth-First Search): Only works for unweighted graphs or graphs where all edges have equal weight. Our campus paths have different lengths.")
    add_bullet(doc, "Bellman-Ford: Handles negative weights but runs in O(VE) — slower and unnecessary when all weights are non-negative.")
    add_bullet(doc, "A* (A-star): An extension of Dijkstra that uses heuristics (e.g., straight-line distance to goal). Useful for large maps; our small campus graph does not require it.")
    add_bullet(doc, "Floyd-Warshall: Computes all-pairs shortest paths in O(V³). Overkill when we only need one source-destination pair.")

    add_para(
        doc,
        "Dijkstra's Algorithm with a priority queue is the optimal choice for this application: "
        "correct, efficient, and well-suited to weighted graphs with non-negative edge weights."
    )

    # 9. Summary
    add_heading(doc, "9. Summary", level=1)
    add_para(
        doc,
        "The Campus Route Planner demonstrates the practical application of fundamental data "
        "structures and algorithms:"
    )
    add_bullet(doc, "Graph (adjacency list) — models the campus network")
    add_bullet(doc, "Priority queue (binary min-heap) — efficiently selects the next vertex to process")
    add_bullet(doc, "Hash maps — store distances and predecessors for O(1) lookup")
    add_bullet(doc, "Set — track blocked edges for road closure simulation")
    add_bullet(doc, "Dijkstra's Algorithm — computes the shortest path with O((V + E) log V) time complexity")

    add_para(
        doc,
        "Together, these components form a complete solution that mirrors how real navigation "
        "systems (such as Google Maps) compute routes, scaled down to a campus environment."
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Campus Route Planner — Data Structures and Algorithms Project")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    out_path = "Campus_Route_Planner_Algorithm_Explanation.docx"
    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
